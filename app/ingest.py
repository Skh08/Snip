"""Structure-preserving DOCX ingestion for СНиП documents."""
from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document

from .models import Chunk

# Chapter headings are e.g. "4. НАРУЖНЫЕ ГАЗОПРОВОДЫ"; do not treat "4.17." as a heading.
SECTION_RE = re.compile(r"^\d+\.\s+\D.+")
PARAGRAPH_RE = re.compile(r"^(\d+(?:\.\d+)+)\.?\s*(.*)$")
APPENDIX_RE = re.compile(r"^(?:Приложение|ПРИЛОЖЕНИЕ)\b", re.IGNORECASE)
TABLE_RE = re.compile(r"^(?:Таблица|ТАБЛИЦА)\s+(\d+)", re.IGNORECASE)


def _clean(value: str) -> str:
    return re.sub(r"\s+", " ", value).strip()


def _label(section: str | None, paragraph: str | None, fallback: str) -> str:
    if paragraph:
        return f"п. {paragraph}"
    if section:
        return section
    return fallback


def parse_docx(path: Path) -> list[Chunk]:
    document = Document(path)
    chunks: list[Chunk] = []
    section: str | None = None
    active_table: str | None = None
    ordinal = 0

    for raw_paragraph in document.paragraphs:
        text = _clean(raw_paragraph.text)
        if not text:
            continue
        if SECTION_RE.match(text) or APPENDIX_RE.match(text):
            section = text
            active_table = None
        table_match = TABLE_RE.match(text)
        if table_match:
            active_table = f"Таблица {table_match.group(1)}"
        match = PARAGRAPH_RE.match(text)
        paragraph = match.group(1) if match else None
        # A numbered regulatory paragraph marks the end of a text-formatted table.
        if paragraph and not table_match:
            active_table = None
        ordinal += 1
        chunks.append(Chunk(
            id=f"p-{ordinal}", text=text, section=section, paragraph=paragraph,
            source_label=(active_table if active_table and not paragraph else _label(section, paragraph, f"абзац {ordinal}")),
            kind="table_text" if active_table and not paragraph else "paragraph", ordinal=ordinal,
        ))

    for table_index, table in enumerate(document.tables, start=1):
        rows = []
        for row in table.rows:
            cells = [_clean(cell.text) for cell in row.cells]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            ordinal += 1
            chunks.append(Chunk(
                id=f"t-{table_index}", text="\n".join(rows), section=section,
                source_label=f"Таблица {table_index}" + (f" — {section}" if section else ""),
                kind="table", ordinal=ordinal,
            ))
    return chunks


def write_knowledge_base(source: Path, destination: Path) -> int:
    chunks = parse_docx(source)
    destination.parent.mkdir(parents=True, exist_ok=True)
    destination.write_text(json.dumps([chunk.model_dump() for chunk in chunks], ensure_ascii=False, indent=2), encoding="utf-8")
    return len(chunks)
